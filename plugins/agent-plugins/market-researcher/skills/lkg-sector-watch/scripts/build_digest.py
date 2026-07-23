#!/usr/bin/env python3
"""
Render the LKG Portfolio Sector Digest to .docx and re-enforce the flag gate.

The agent nominates flags; this script is the deterministic gate. Any flag
missing a required field or carrying an out-of-enum value is demoted to the
Watchlist with the failure named — it never renders as a flag.

Usage: python build_digest.py digest.json [output.docx]
Requires: python-docx  (pip install python-docx)

DESIGN RULES (the "why", so future edits stay on-style)
-------------------------------------------------------
1. Headings are navy bands (white uppercase text on NAVY fill). Every section
   gets a one-line italic grey subtitle explaining what the section is for.
2. Colour is semantic only: RED = high severity / risk / draft warnings,
   AMBER = med severity, GREEN = opportunity / up-delta, GREY = low severity /
   metadata. Body text stays near-black. No decorative colour.
3. Flag cards are compact 4-line single-column tables:
     [shaded band]  SEV · AUDIENCE  Headline (bold)
     Polarity · mechanism            (one metadata line)
     Why this matters: ...           (literal label — matches the brief)
     Source: <hyperlinked label>
   Hairline outer border, no inside borders. Band fill by severity.
4. No raw URLs in the body — sources are short hyperlinked labels with a
   trailing " ↗". Full URLs live only in the Appendix (new page, small grey),
   de-duplicated and tagged with the sections that cite them.
5. "Why this matters" appears with those literal words everywhere an
   implication is stated (flag cards and sector bullets).
6. Watchlist = adjacent items at a glance: visually demoted (smaller, grey,
   status tag first). Sector detail = core section, full-size text, bold fact
   line + indented "Why this matters:" line.
7. Macro dashboard: navy header row, zebra striping, ▲/▼ direction column
   (green/red). No per-row sector column in a single-sector run.
8. Reviewer decisions: ☐ Promote / ☐ Reject / ☐ Amend checkboxes per row,
   truncated flag references (~9 words — §1 carries the full text), red italic
   "no distribution until completed" line above, sign-off below.
9. Assumptions & audit trail: small grey text — present but visually
   subordinate to the flags.
"""
import json
import re
import sys
from pathlib import Path

REQUIRED = ["headline", "source_url", "polarity", "mechanism",
            "implication", "owner", "severity"]
HEADLINE_MAX = 140  # the schema's "one line", made checkable
MECHANISMS = {  # the frozen taxonomy from references/sectors/ — no extensions mid-run
    "housing-turnover", "rates-sentiment", "input-costs", "seasonality",
    "franchise-risk", "ma-distress", "competitor-move",
}
ENUMS = {
    "polarity": {"risk", "opportunity", "watch"},
    "owner": {"GM", "board", "QLC"},
    "severity": {"low", "med", "high"},
    "mechanism": MECHANISMS,
}
SEV_ORDER = {"high": 0, "med": 1, "low": 2}
OWNER_ORDER = {"board": 0, "QLC": 1, "GM": 2}

PRIMARY_LABELS = {  # originators get clean names; everything else shows its bare domain
    "cotality.com.au": "Cotality", "abs.gov.au": "ABS", "accc.gov.au": "ACCC",
    "westpaciq.com.au": "Westpac IQ", "rba.gov.au": "RBA", "asx.com.au": "ASX",
}
MONTHS = {"01": "January", "02": "February", "03": "March", "04": "April",
          "05": "May", "06": "June", "07": "July", "08": "August",
          "09": "September", "10": "October", "11": "November",
          "12": "December"}


def gate(flags):
    """Split nominated flags into (passed, demoted-with-reason)."""
    passed, demoted = [], []
    for f in flags:
        missing = [k for k in REQUIRED if not str(f.get(k) or "").strip()]
        bad = [k for k, allowed in ENUMS.items()
               if f.get(k) and f[k] not in allowed]
        head = str(f.get("headline", "")).strip()
        over = len(head) > HEADLINE_MAX
        multiline = "\n" in head
        if missing or bad or over or multiline:
            reason = "failed gate: " + "; ".join(
                (["missing " + ", ".join(missing)] if missing else []) +
                [f"invalid {k}={f[k]!r}" for k in bad] +
                ([f"headline over {HEADLINE_MAX} chars"] if over else []) +
                (["headline not one line"] if multiline else []))
            demoted.append({"headline": str(f.get("headline") or
                                            "(no headline)"),
                            "source_url": str(f.get("source_url") or ""),
                            "reason": reason})
        else:
            passed.append(f)
    passed.sort(key=lambda f: (SEV_ORDER[f["severity"]],
                               OWNER_ORDER[f["owner"]]))
    return passed, demoted


def source_label(url):
    host = re.sub(r"^https?://(www\.)?", "", url or "").split("/")[0]
    for dom, label in PRIMARY_LABELS.items():
        if host == dom or host.endswith("." + dom):
            return label
    return host or "source"


def direction(delta):
    d = delta.strip().lower()
    if d.startswith(("+", "up")) or "rose" in d[:24]:
        return "up"
    if d.startswith(("-", "−", "down", "below")) or "fell" in d[:24]:
        return "down"
    return ""


def pretty_date(iso):
    try:
        y, m, d = iso.split("-")
        return f"{int(d)} {MONTHS[m]} {y}"
    except (ValueError, KeyError):
        return iso


def split_reason(reason):
    """'failed gate: details' -> ('FAILED GATE', 'details'); free text -> WATCH."""
    m = re.match(r"^([^:]{3,40}):\s*(.*)$", reason or "")
    if m:
        return m.group(1).upper(), m.group(2)
    return "WATCH", reason or ""


def main():
    if len(sys.argv) < 2:
        sys.exit(__doc__.strip().split("\n\nDESIGN RULES")[0])
    data = json.loads(Path(sys.argv[1]).read_text(encoding="utf-8"))
    out = sys.argv[2] if len(sys.argv) > 2 else \
        f"LKG-Sector-Digest-{data['date']}.docx"

    flags, demoted = gate(data.get("flags", []))
    watchlist = demoted + list(data.get("watchlist", []))

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        sys.exit("python-docx not installed (pip install python-docx). "
                 "Gate result: %d flag(s) passed, %d demoted." %
                 (len(flags), len(demoted)))

    NAVY, AMBER, AMBERBG = "1F3A5F", "B45309", "FDF3E3"
    GREY, GREYBG, DARKGREY = "6B7280", "F3F4F6", "374151"
    RED, REDBG = "B91C1C", "FBEAEA"
    GREEN, LIGHT, DARK = "15803D", "D1D5DB", "111827"
    LINKBLUE = "1D4ED8"
    FONT = "Calibri"
    PAGE = 9706  # A4 width (11906) minus the 1100-twip side margins below
    SEV_TXT = {"high": RED, "med": AMBER, "low": GREY}
    SEV_BG = {"high": REDBG, "med": AMBERBG, "low": GREYBG}
    SEV_LEGEND = {"high": "act this week", "med": "act on this quarter",
                  "low": "context / monitor"}
    AUD = {"board": "Board", "GM": "GM", "QLC": "QLC"}
    POL = {"opportunity": GREEN, "watch": GREY, "risk": RED}

    def rgb(hexstr):
        return RGBColor.from_string(hexstr)

    def run(par, text, size=10, color=DARK, bold=False, italic=False,
            caps=False):
        r = par.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.color.rgb = rgb(color)
        r.font.bold = bold
        r.font.italic = italic
        if caps:
            r.font.all_caps = True
        return r

    def spacing(par, before=0, after=0):
        par.paragraph_format.space_before = Pt(before)
        par.paragraph_format.space_after = Pt(after)

    def hyperlink(par, label, url, size=8.5):
        r_id = par.part.relate_to(url, RT.HYPERLINK, is_external=True)
        hl = OxmlElement("w:hyperlink")
        hl.set(qn("r:id"), r_id)
        r, rpr = OxmlElement("w:r"), OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), FONT)
        fonts.set(qn("w:hAnsi"), FONT)
        rpr.append(fonts)
        rpr.append(OxmlElement("w:b"))
        for tag, val in (("w:color", LINKBLUE), ("w:sz", str(int(size * 2))),
                         ("w:u", "single")):
            el = OxmlElement(tag)
            el.set(qn("w:val"), val)
            rpr.append(el)
        t = OxmlElement("w:t")
        t.text = label + " ↗"
        t.set(qn("xml:space"), "preserve")
        r.append(rpr)
        r.append(t)
        hl.append(r)
        par._p.append(hl)

    def shade_par(par, fill):
        # shd must precede spacing/ind/jc in pPr per CT_PPr child order
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill)
        par._p.get_or_add_pPr().insert_element_before(
            shd, "w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr",
            "w:sectPr", "w:pPrChange")

    def shade_cell(cell, fill):
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill)
        cell._tc.get_or_add_tcPr().insert_element_before(
            shd, "w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText",
            "w:vAlign", "w:hideMark", "w:tcPrChange")

    def cell_margins(cell, top=30, bottom=30, left=140, right=140):
        mar = OxmlElement("w:tcMar")
        for tag, val in (("top", top), ("left", left), ("bottom", bottom),
                         ("right", right)):
            el = OxmlElement(f"w:{tag}")
            el.set(qn("w:w"), str(val))
            el.set(qn("w:type"), "dxa")
            mar.append(el)
        cell._tc.get_or_add_tcPr().insert_element_before(
            mar, "w:textDirection", "w:tcFitText", "w:vAlign", "w:hideMark",
            "w:tcPrChange")

    def table_borders(table, inside=False):
        borders = OxmlElement("w:tblBorders")
        for tag in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{tag}")
            if tag.startswith("inside") and not inside:
                el.set(qn("w:val"), "none")
                el.set(qn("w:sz"), "0")
            else:
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), "4")
                el.set(qn("w:color"), LIGHT)
            borders.append(el)
        table._tbl.tblPr.insert_element_before(
            borders, "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook",
            "w:tblCaption", "w:tblDescription", "w:tblPrChange")

    def fixed_widths(table, widths):
        table.autofit = False  # python-docx emits <w:tblLayout w:type="fixed"/>
        for row in table.rows:
            for i, w in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = Twips(w)

    def cellp(cell):
        return cell.paragraphs[0]

    def heading_band(num, text):
        p = doc.add_paragraph()
        spacing(p, before=18, after=3)
        shade_par(p, NAVY)
        run(p, f"  {num}   ", size=13.5, color="C8D3E4", bold=True)
        run(p, text.upper(), size=13.5, color="FFFFFF", bold=True)

    def subtitle(text):
        p = doc.add_paragraph()
        spacing(p, after=7)
        run(p, text, size=9, color=GREY, italic=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Twips(11906)   # A4
    section.page_height = Twips(16838)
    section.top_margin = Twips(900)
    section.bottom_margin = Twips(900)
    section.left_margin = Twips(1100)
    section.right_margin = Twips(1100)
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = Pt(10)

    # ------------------------------------------------------- title block ----
    sectors = data.get("sectors", [])
    macro = data.get("macro", [])
    p = doc.add_paragraph()
    spacing(p, after=2)
    run(p, "LKG Portfolio Sector Digest", size=20, bold=True, color=NAVY)
    p = doc.add_paragraph()
    spacing(p, after=3)
    run(p, "  ·  ".join([s["name"] for s in sectors] or ["LKG portfolio"]) +
        f"  ·  {pretty_date(data['date'])}", size=10.5, color=GREY)
    p = doc.add_paragraph()
    spacing(p, after=2)
    run(p, "DRAFT — NOT FOR DISTRIBUTION  ", size=9.5, bold=True, color=RED)
    run(p, f"·  Run {data.get('run_time', '')}  ·  Sources scanned: "
           f"{data.get('sources_scanned', '?')}  ·  For review by "
           f"{data.get('reviewer', '[name]')}", size=9.5, color=GREY)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), NAVY)
    borders.append(bottom)
    p._p.get_or_add_pPr().insert_element_before(
        borders, "w:shd", "w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr",
        "w:sectPr", "w:pPrChange")

    # ----------------------------------------------- executive summary -----
    heading_band("ES", "Executive summary")
    subtitle("The whole digest on one page — narrative first, then the "
             "numbers. The outline is compiled from the gated content below "
             "and cannot drift from it.")
    for b in data.get("executive_summary", []):
        p = doc.add_paragraph(style="List Bullet")
        spacing(p, after=3)
        run(p, b, size=9.5)
    outline = []
    if flags:
        owners = [f["owner"] for f in flags]
        sevs = [f["severity"] for f in flags]
        aud_bits = ", ".join(x for x in (
            f"{owners.count('board')} Board" if "board" in owners else "",
            f"{owners.count('QLC')} QLC" if "QLC" in owners else "",
            f"{owners.count('GM')} GM" if "GM" in owners else "") if x)
        sev_bits = ", ".join(f"{sevs.count(s)} {s}"
                             for s in ("high", "med", "low") if s in sevs)
        outline.append(f"§1 — {len(flags)} item(s) flagged ({aud_bits}; "
                       f"{sev_bits}). Top item: {flags[0]['headline']}.")
    else:
        outline.append("§1 — no items passed the flag gate this run.")
    if macro:
        ups = sum(1 for m in macro
                  if direction(m.get("delta", "")) == "up")
        downs = sum(1 for m in macro
                    if direction(m.get("delta", "")) == "down")
        outline.append(f"§2 — {len(macro)} macro indicator(s): "
                       f"{downs} moving down, {ups} up.")
    else:
        outline.append("§2 — no macro deltas this run (deltas only).")
    n_items = sum(len(s.get("items", [])) for s in sectors)
    if len(sectors) == 1:
        outline.append(f"§3 — {n_items} item(s) for {sectors[0]['name']}.")
    else:
        outline.append(f"§3 — {n_items} item(s) across "
                       f"{len(sectors)} sector(s).")
    outline.append(f"§4 — {len(watchlist)} watchlist item(s)" +
                   (f", {len(demoted)} demoted by the flag gate with the "
                    f"failure named" if demoted else "") + ".")
    outline.append("§6 — action required: the named reviewer completes one "
                   "Promote / Reject / Amend row per flag; nothing is "
                   "distributed until then.")
    for text in outline:
        p = doc.add_paragraph(style="List Bullet")
        spacing(p, after=2)
        run(p, text, size=8.5, color=DARKGREY)

    # ------------------------------------------------------ 1 · flags ------
    heading_band("1", "Flagged items")
    subtitle("Everything that cleared the deterministic flag gate this run — "
             "nominations for the named reviewer, not decisions.")
    p = doc.add_paragraph()
    spacing(p, after=7)
    wrote_legend = False
    for sev in ("high", "med", "low"):
        if any(f["severity"] == sev for f in flags):
            run(p, sev.upper(), size=8.5, bold=True, color=SEV_TXT[sev])
            run(p, f" {SEV_LEGEND[sev]}    ", size=8.5, color=GREY)
            wrote_legend = True
    run(p, ("·    " if wrote_legend else "") +
           "Audience: Board, GM or QLC    ·    All flags pending "
           "reviewer decision (§6)", size=8.5, color=GREY)
    if not flags:
        p = doc.add_paragraph()
        run(p, "No items passed the flag gate this run.", size=10)
    for f in flags:
        sev = f["severity"]
        t = doc.add_table(rows=4, cols=1)
        table_borders(t)
        fixed_widths(t, [PAGE])
        c = t.rows[0].cells[0]
        shade_cell(c, SEV_BG[sev])
        cell_margins(c, top=80, bottom=80)
        p = cellp(c)
        run(p, f"{sev.upper()} · {AUD[f['owner']].upper()}   ", size=8.5,
            color=SEV_TXT[sev], bold=True)
        run(p, f["headline"], size=10.5, bold=True)
        c = t.rows[1].cells[0]
        cell_margins(c, top=70, bottom=20)
        p = cellp(c)
        run(p, f["polarity"].capitalize(), size=9, bold=True,
            color=POL[f["polarity"]])
        run(p, f"  ·  {f['mechanism']}", size=9, color=GREY)
        c = t.rows[2].cells[0]
        cell_margins(c)
        p = cellp(c)
        run(p, "Why this matters: ", size=9.5, bold=True, color=NAVY)
        run(p, f["implication"], size=9.5)
        c = t.rows[3].cells[0]
        cell_margins(c, top=20, bottom=80)
        p = cellp(c)
        run(p, "Source: ", size=8.5, color=GREY)
        hyperlink(p, source_label(f["source_url"]), f["source_url"])
        gap = doc.add_paragraph()
        spacing(gap, after=7)

    # ------------------------------------------------------ 2 · macro ------
    heading_band("2", "Macro dashboard")
    subtitle("Latest prints for the demand indicators the sector trades on.")
    if macro:
        touched = {s for m in macro for s in m.get("sectors_touched", [])}
        multi = len(touched) > 1
        widths = [3900, 3306, 1900, 600] if multi else [4750, 4356, 600]
        heads = ("Indicator", "Delta") + \
            (("Sectors",) if multi else ()) + ("",)
        t = doc.add_table(rows=1 + len(macro), cols=len(heads))
        table_borders(t)
        fixed_widths(t, widths)
        for i, h in enumerate(heads):
            c = t.rows[0].cells[i]
            shade_cell(c, NAVY)
            cell_margins(c, top=60, bottom=60, left=100, right=100)
            run(cellp(c), h, size=9, bold=True, color="FFFFFF", caps=True)
        for ri, m in enumerate(macro, start=1):
            cells = t.rows[ri].cells
            for c in cells:
                if ri % 2 == 0:
                    shade_cell(c, GREYBG)
                cell_margins(c, top=50, bottom=50, left=100, right=100)
            run(cellp(cells[0]), m.get("indicator", ""), size=9.5)
            run(cellp(cells[1]), m.get("delta", ""), size=9.5)
            i = 2
            if multi:
                run(cellp(cells[2]),
                    ", ".join(m.get("sectors_touched", [])), size=9.5)
                i = 3
            p = cellp(cells[i])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            d = direction(m.get("delta", ""))
            if d:
                run(p, "▲" if d == "up" else "▼", size=9.5,
                    bold=True, color=GREEN if d == "up" else RED)
    else:
        p = doc.add_paragraph()
        run(p, "No macro deltas this run (deltas only — non-events are not "
               "reported).", size=9.5, color=DARKGREY)

    # ----------------------------------------------- 3 · sector detail -----
    active = [s for s in sectors if s.get("status") != "stub"]
    band = "Sector detail"
    if len(active) == 1 and len(sectors) == 1:
        band += f" — {active[0]['name']}"
    heading_band("3", band)
    subtitle("Core section — the active sector in full. Each item states the "
             "fact, why it matters to the portfolio, and its source.")
    for s in sectors:
        if len(sectors) > 1 or s.get("status") == "stub":
            p = doc.add_paragraph()
            spacing(p, before=4, after=3)
            run(p, f"{s['name']} ({s.get('status', 'active')})", size=10.5,
                bold=True, color=NAVY)
        if s.get("status") == "stub":
            p = doc.add_paragraph()
            run(p, s.get("note", "not yet configured — sectors onboarded "
                                 "deliberately"), size=9, color=GREY,
                italic=True)
            continue
        for it in s.get("items", []):
            p = doc.add_paragraph(style="List Bullet")
            spacing(p, after=1)
            aud = AUD.get(it.get("audience", ""), it.get("audience", "?"))
            run(p, f"[{aud}] ", size=9.5, bold=True,
                color=AMBER if aud == "Board" else GREY)
            run(p, it["headline"], size=9.5, bold=True)
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Twips(360)
            spacing(p2, after=6.5)
            run(p2, "Why this matters: ", size=9.5, bold=True, color=NAVY)
            run(p2, (it.get("implication", "") or "").rstrip() + "  ",
                size=9.5)
            if it.get("source_url"):
                run(p2, "Source: ", size=8.5, color=GREY)
                hyperlink(p2, source_label(it["source_url"]),
                          it["source_url"])

    # -------------------------------------------------- 4 · watchlist ------
    heading_band("4", "Watchlist — adjacent items at a glance")
    subtitle("Comps, pending events and gate demotions monitored for "
             "read-through. Nothing here is actionable this run — items "
             "graduate to sections 1–3 when they become material.")
    for w in watchlist:
        p = doc.add_paragraph(style="List Bullet")
        spacing(p, after=4.5)
        status, detail = split_reason(w.get("reason", ""))
        run(p, status + "  ", size=7.5, bold=True, color=GREY)
        run(p, w["headline"] + (f" — {detail}" if detail else "") + "  ",
            size=8.5, color=DARKGREY)
        if w.get("source_url"):
            hyperlink(p, source_label(w["source_url"]), w["source_url"],
                      size=7.5)
    if not watchlist:
        p = doc.add_paragraph()
        run(p, "Empty.", size=9, color=GREY)

    # ------------------------------------------------ 5 · assumptions ------
    heading_band("5", "Assumptions & audit trail")
    subtitle("How this run was produced — scope, source caps, blind spots, "
             "and items dropped at the verification gate.")
    for a in data.get("assumptions", []):
        p = doc.add_paragraph(style="List Bullet")
        spacing(p, after=3.5)
        run(p, a, size=8.5, color=GREY)

    # ------------------------------------------- 6 · reviewer decisions ----
    heading_band("6", "Reviewer decisions")
    subtitle("The human approval point — the agent stops here; only the "
             "named reviewer completes this section.")
    p = doc.add_paragraph()
    spacing(p, after=5)
    run(p, "No item in section 1 is distributed until its row below is "
           "completed.", size=9.5, italic=True, color=RED)
    t = doc.add_table(rows=1 + max(len(flags), 1), cols=3)
    table_borders(t, inside=True)
    fixed_widths(t, [4500, 2206, 3000])
    for i, h in enumerate(("Flag", "Decision", "Reviewer comment")):
        c = t.rows[0].cells[i]
        shade_cell(c, NAVY)
        cell_margins(c, top=60, bottom=60, left=100, right=100)
        run(cellp(c), h, size=9, bold=True, color="FFFFFF", caps=True)
    if not flags:
        c = t.rows[1].cells[0]
        cell_margins(c, top=70, bottom=70, left=100, right=100)
        run(cellp(c), "(no flags passed the gate this run)", size=9,
            color=GREY)
    for ri, f in enumerate(flags, start=1):
        words = f["headline"].split()
        short = " ".join(words[:9]) + ("…" if len(words) > 9 else "")
        c = t.rows[ri].cells[0]
        cell_margins(c, top=70, bottom=70, left=100, right=100)
        p = cellp(c)
        run(p, f"[{f['severity'].upper()} · {AUD[f['owner']]}]  ", size=8.5,
            bold=True, color=SEV_TXT[f["severity"]])
        run(p, short, size=9)
        c = t.rows[ri].cells[1]
        cell_margins(c, top=70, bottom=70, left=100, right=100)
        p = cellp(c)
        run(p, "☐ Promote", size=9)
        for opt in ("☐ Reject", "☐ Amend"):
            p = c.add_paragraph()
            spacing(p)
            run(p, opt, size=9)
        cell_margins(t.rows[ri].cells[2], top=70, bottom=70, left=100,
                     right=100)
    p = doc.add_paragraph()
    spacing(p, before=12)
    run(p, "Reviewed by:  ______________________________        "
           "Date:  ______________", size=10)

    # --------------------------------------------- appendix · sources ------
    # The one place full URLs appear — every body link resolves to an entry
    # here. De-duplicated by URL, first-appearance order, tagged by section.
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    heading_band("A", "Appendix — Sources")
    subtitle("Every source cited in this digest, with full URLs for audit. "
             "Body links resolve to the entries below.")
    seen = {}

    def collect(entries, tag, key="source_url"):
        for e in entries:
            u = e.get(key)
            if not u:
                continue
            if u not in seen:
                seen[u] = {"label": source_label(u), "sections": []}
            if tag not in seen[u]["sections"]:
                seen[u]["sections"].append(tag)

    collect(flags, "§1")
    collect(macro, "§2")
    for s in sectors:
        collect(s.get("items", []), "§3")
    collect(watchlist, "§4")
    for i, (url, meta) in enumerate(seen.items(), start=1):
        p = doc.add_paragraph()
        spacing(p, after=1)
        run(p, f"{i}.  ", size=9, bold=True, color=NAVY)
        run(p, meta["label"], size=9, bold=True)
        run(p, "   (cited in " + ", ".join(meta["sections"]) + ")", size=8,
            color=GREY)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Twips(280)
        spacing(p2, after=6)
        run(p2, url, size=8, color=GREY)

    doc.save(out)
    print(f"wrote {out}: {len(flags)} flag(s), {len(demoted)} demoted to "
          f"watchlist by gate, {len(data.get('sectors', []))} sector(s)")


if __name__ == "__main__":
    main()
