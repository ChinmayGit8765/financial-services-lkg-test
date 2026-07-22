#!/usr/bin/env python3
"""
Render the LKG Portfolio Sector Digest to .docx and re-enforce the flag gate.

The agent nominates flags; this script is the deterministic gate. Any flag
missing a required field or carrying an out-of-enum value is demoted to the
Watchlist with the failure named — it never renders as a flag.

Usage: python build_digest.py digest.json [output.docx]
Requires: python-docx  (pip install python-docx)
"""
import json
import sys
from pathlib import Path

REQUIRED = ["headline", "source_url", "polarity", "mechanism",
            "implication", "owner", "severity"]
ENUMS = {
    "polarity": {"risk", "opportunity", "watch"},
    "owner": {"GM", "board", "QLC"},
    "severity": {"low", "med", "high"},
}
SEV_ORDER = {"high": 0, "med": 1, "low": 2}
OWNER_ORDER = {"board": 0, "QLC": 1, "GM": 2}


def gate(flags):
    """Split nominated flags into (passed, demoted-with-reason)."""
    passed, demoted = [], []
    for f in flags:
        missing = [k for k in REQUIRED if not str(f.get(k, "")).strip()]
        bad = [k for k, allowed in ENUMS.items()
               if f.get(k) and f[k] not in allowed]
        if missing or bad:
            reason = "failed gate: " + "; ".join(
                (["missing " + ", ".join(missing)] if missing else []) +
                [f"invalid {k}={f[k]!r}" for k in bad])
            demoted.append({"headline": f.get("headline", "(no headline)"),
                            "source_url": f.get("source_url", ""),
                            "reason": reason})
        else:
            passed.append(f)
    passed.sort(key=lambda f: (SEV_ORDER[f["severity"]],
                               OWNER_ORDER[f["owner"]]))
    return passed, demoted


def main():
    if len(sys.argv) < 2:
        sys.exit(__doc__.strip())
    data = json.loads(Path(sys.argv[1]).read_text(encoding="utf-8"))
    out = sys.argv[2] if len(sys.argv) > 2 else \
        f"LKG-Sector-Digest-{data['date']}.docx"

    flags, demoted = gate(data.get("flags", []))
    watchlist = demoted + list(data.get("watchlist", []))

    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        sys.exit("python-docx not installed (pip install python-docx). "
                 "Gate result: %d flag(s) passed, %d demoted." %
                 (len(flags), len(demoted)))

    doc = Document()
    doc.add_heading(f"LKG Portfolio Sector Digest — {data['date']}", 0)
    hdr = doc.add_paragraph()
    hdr.add_run(
        f"Run: {data.get('run_time', '')} · Sources scanned: "
        f"{data.get('sources_scanned', '?')} · ").italic = True
    d = hdr.add_run(f"DRAFT — for review by {data.get('reviewer', '[name]')}")
    d.bold = True

    doc.add_heading("1. Flagged items", 1)
    if not flags:
        doc.add_paragraph("No items passed the flag gate this run.")
    for f in flags:
        doc.add_heading(
            f"[{f['severity'].upper()} · {f['owner']}] {f['headline']}", 2)
        t = doc.add_table(rows=0, cols=2)
        for k in ("polarity", "mechanism", "implication",
                  "sector", "source_url"):
            if f.get(k):
                row = t.add_row().cells
                row[0].text, row[1].text = k, str(f[k])

    doc.add_heading("2. Macro dashboard", 1)
    if data.get("macro"):
        t = doc.add_table(rows=1, cols=3)
        for i, h in enumerate(("Indicator", "Delta", "Sectors touched")):
            t.rows[0].cells[i].text = h
        for m in data["macro"]:
            row = t.add_row().cells
            row[0].text = m.get("indicator", "")
            row[1].text = m.get("delta", "")
            row[2].text = ", ".join(m.get("sectors_touched", []))
    else:
        doc.add_paragraph("No macro deltas this run (deltas only — "
                          "non-events are not reported).")

    doc.add_heading("3. Sector sections", 1)
    for s in data.get("sectors", []):
        doc.add_heading(f"{s['name']} ({s.get('status', 'active')})", 2)
        if s.get("status") == "stub":
            doc.add_paragraph(s.get(
                "note", "not yet configured — sectors onboarded deliberately"))
            continue
        for it in s.get("items", []):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{it.get('audience', '?')}] ").bold = True
            p.add_run(f"{it['headline']} — {it.get('implication', '')} ")
            p.add_run(f"({it.get('source_url', 'no source')})").italic = True

    doc.add_heading("4. Watchlist", 1)
    for w in watchlist:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(w["headline"])
        p.add_run(f" — {w.get('reason', '')} "
                  f"({w.get('source_url', 'no source')})").italic = True
    if not watchlist:
        doc.add_paragraph("Empty.")

    doc.add_heading("5. Assumptions appendix", 1)
    for a in data.get("assumptions", []):
        doc.add_paragraph(a, style="List Bullet")

    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.size is None:
                r.font.size = Pt(10)
    doc.save(out)
    print(f"wrote {out}: {len(flags)} flag(s), {len(demoted)} demoted to "
          f"watchlist by gate, {len(data.get('sectors', []))} sector(s)")


if __name__ == "__main__":
    main()
